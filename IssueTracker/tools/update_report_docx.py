from pathlib import Path
import shutil
import tempfile
import textwrap

import fitz
from docx import Document
from docx.shared import Inches


SOURCE_PDF = Path(r"C:\Users\STAR LAPTOPS\Downloads\Student Issue Tracker (1).pdf")
OUT_DIR = Path(r"D:\IssueTracker")
UPDATED_PDF = OUT_DIR / "Student_Issue_Tracker_Updated_Workflow_Final_v2.pdf"
OUT_DOCX = OUT_DIR / "Student_Issue_Tracker_Updated_Workflow_Final_v2.docx"


def redact(page, rect, text="", fontsize=11, fontname="Times-Roman", color=(0, 0, 0)):
    text = text.replace("•", "-").replace("—", "-").replace("–", "-")
    page.add_redact_annot(fitz.Rect(rect), fill=(1, 1, 1))
    page.apply_redactions()
    if text:
        page.insert_textbox(
            fitz.Rect(rect),
            textwrap.dedent(text).strip(),
            fontsize=fontsize,
            fontname=fontname,
            color=color,
            align=fitz.TEXT_ALIGN_LEFT,
        )


def replace_area(page, rect, text, fontsize=11):
    redact(page, rect, text=text, fontsize=fontsize)


def update_pdf():
    doc = fitz.open(SOURCE_PDF)

    # Page 5: project brief values.
    p = doc[4]
    redact(p, (258, 526, 548, 566), "JavaScript (Next.js/React), CSS, SQL,\nFramer Motion")
    redact(p, (258, 572, 548, 596), "Supabase PostgreSQL (Development/Production)")
    redact(
        p,
        (258, 611, 552, 681),
        "Next.js API Routes, Supabase Auth/Storage,\nRecharts, bcryptjs, Botpress, VS Code,\nPostman, Git/GitHub",
    )

    # Page 6: abstract.
    replace_area(
        doc[5],
        (72, 154, 548, 382),
        """
        The AI-Enhanced Student Issue Tracker is a web-based complaint management platform
        designed to streamline issue handling in university departments.

        • It supports multiple user roles including Student, Faculty Member, HOD, DSA, and
        Supervisor, ensuring a structured workflow from submission to closure.
        • The system includes AI-assisted category suggestion, severity/priority estimation,
        duplicate-aware submission support, attachment handling, notifications, analytics,
        activity logs, weekly reporting, and chatbot assistance.
        • Complaints are routed by category and department: Academic, Facilities, Behavior-
        related, and Other issues normally route to HOD, while Administrative issues route
        to DSA. HOD/DSA can review, reject, or assign work to faculty; faculty resolve cases;
        DSA/Supervisor close finalized complaints; students can rate the resolution.

        Developed using Next.js, React, Supabase Auth, Supabase PostgreSQL, and server-side
        Next.js API routes, the system improves transparency, accountability, and resolution
        efficiency.
        """,
        fontsize=10.5,
    )

    # Page 13: objectives.
    replace_area(
        doc[12],
        (88, 130, 545, 306),
        """
        1. To design and implement a web-based AI-enhanced department complaint management
        system using Next.js, React, and Supabase.
        2. To provide secure role-based registration and login for Student, Faculty Member,
        HOD, DSA, and Supervisor.
        3. To suggest complaint category and priority during submission and support duplicate-
        aware complaint handling.
        4. To implement a structured workflow from student submission to HOD/DSA review,
        faculty assignment, resolution, closure, and student feedback.
        5. To provide role-based dashboards, notifications, analytics, activity tracking, and
        weekly reports for transparent decision-making.
        6. To provide chatbot-assisted guidance for student users throughout the complaint
        process.
        """,
        fontsize=10.2,
    )

    # Page 17: architecture overview.
    replace_area(
        doc[16],
        (72, 70, 548, 724),
        """
        Presentation Tier (Next.js/React Frontend): The frontend is implemented with Next.js
        App Router and React components. It provides the landing page, registration/login
        screens, role-aware dashboard, complaint form, complaint cards, analytics panels,
        notifications, Botpress student assistant, and responsive user interface components.
        Client-side helpers call protected API routes and Supabase session handling keeps each
        user inside the dashboard scope allowed for their role.

        Application Tier (Next.js API Routes): Server-side API routes provide the business
        workflow for authentication profile lookup, complaint creation, review, assignment,
        status updates, finalization, ratings, comments, attachments, analytics, activity logs,
        notifications, departments, staff management, and weekly reports. All protected routes
        validate the Supabase session token and then enforce role permissions before running
        the requested action.

        Data Tier (Supabase PostgreSQL): Supabase stores profiles, departments, complaints,
        complaint attachments, comments, notifications, category routes, and activity logs.
        PostgreSQL enums define roles, statuses, categories, priorities, and routed roles.
        Row Level Security policies protect user-owned records, while trusted server-side API
        routes use the service role for cross-role workflows such as staff assignment, analytics,
        and supervisor operations.

        Workflow Layer: Category routing maps Academic, Behavior-related, Facilities, and
        Other complaints to HOD by default, and Administrative complaints to DSA. Students
        submit complaints with optional attachments. HOD/DSA users review and accept/reject
        submitted complaints, then assign accepted work to faculty members. Faculty members
        update assigned complaints and mark them resolved. DSA or Supervisor closes finalized
        complaints, after which students may provide rating and feedback.

        Support Services: The system includes email/notification helpers, analytics summaries,
        activity tracking, weekly reporting, password validation, Botpress chatbot support, and
        reusable UI components for status badges, role badges, timelines, AI panels, and
        complaint cards.
        """,
        fontsize=10.2,
    )

    # Page 23: requirements.
    replace_area(
        doc[22],
        (88, 552, 545, 722),
        """
        • User registration, Supabase session login, and role-based dashboard routing for
        Student, Faculty Member, HOD, DSA, and Supervisor.
        • Complaint submission with title, description, category/priority support, department
        routing, optional anonymous flag, and attachment upload.
        • Automatic routing to HOD or DSA according to category and department rules.
        • HOD/DSA review with accept/reject actions and faculty assignment.
        • Faculty work queue with comments and resolved status updates.
        • DSA/Supervisor finalization, student rating, feedback, and complaint comments.
        • Notifications, activity logs, analytics, weekly reports, department management, and
        chatbot assistance.
        """,
        fontsize=10.3,
    )

    # Page 26: system design and actors.
    replace_area(
        doc[25],
        (72, 108, 545, 384),
        """
        • To define communication pathways between the Next.js/React frontend, server-side
        Next.js API routes, Supabase Auth, Supabase PostgreSQL, storage, notification helpers,
        analytics modules, and chatbot integration.
        • To model system interactions using UML Use Case and Sequence Diagrams while
        keeping the implementation aligned with the current web application.
        • To design a normalized relational database for profiles, departments, complaints,
        attachments, comments, notifications, category routes, and activity logs.

        3.2 Proposed System and its Features
        The system is organized into three logical layers:
        • Layer 1 — Frontend (Next.js/React): landing, login, registration, role dashboards,
        complaint form, AI panel, status badges, role badges, analytics, and Botpress chatbot.
        • Layer 2 — Application/API (Next.js API Routes): Supabase session validation,
        role-based access control, complaint lifecycle actions, staff/department management,
        notifications, analytics, activity feed, and weekly reports.
        • Layer 3 — Data/Storage (Supabase): PostgreSQL tables, Auth users, Row Level
        Security policies, and complaint attachment bucket.
        """,
        fontsize=10.1,
    )
    replace_area(
        doc[25],
        (72, 458, 545, 690),
        """
        Figure 3.1 presents the complete Use Case Diagram for the system. The system boundary
        encloses all use cases. Five actors are considered in the updated implementation:
        Student, Faculty Member, HOD, DSA, and Supervisor. Student use cases include
        Register/Login, Submit Complaint, Upload Attachment, View Status, Add Comment, Rate
        Resolution, and Use Chatbot Assistant. HOD and DSA use cases include View Routed
        Complaints, Accept/Reject Complaint, Assign Faculty, Monitor Analytics, and Receive
        Alerts. Faculty Member use cases include View Assigned Complaints, Add Comments,
        and Mark Resolved. Supervisor use cases include Full Oversight, Department and Staff
        Management, Activity Logs, Analytics, Weekly Reports, Assignment Support, and final
        complaint closure. Notifications are triggered at key workflow events.
        """,
        fontsize=10.1,
    )

    # Page 35: tools.
    replace_area(
        doc[34],
        (76, 218, 548, 586),
        """
        Frontend Framework      Next.js App Router, React, Framer Motion
        UI Components           lucide-react, reusable dashboard components
        Backend Framework       Next.js Server/API Routes
        Authentication          Supabase Auth session tokens
        Database                Supabase PostgreSQL
        Storage                 Supabase Storage complaint-attachments bucket
        Analytics/Charts        Recharts and server-side aggregate API routes
        Password Security       bcryptjs and validation helpers
        Chatbot                 Botpress Cloud Webchat
        Email/Notifications     Next.js notification/email helper modules
        Version Control         Git / GitHub
        IDE / Tools             VS Code, Postman
        """,
        fontsize=10.7,
    )

    # Page 36: implementation and RBAC.
    replace_area(
        doc[35],
        (72, 240, 545, 565),
        """
        4.4.1 Web Portal (Frontend — Next.js/React)
        The frontend is organized around the Next.js App Router with a landing page, login and
        registration screens, and a unified role-aware dashboard. Reusable components render
        complaint cards, complaint forms, role badges, status badges, AI suggestions, timelines,
        analytics, chatbot support, and department management controls.

        4.4.2 Backend (Server-Side — Next.js API Routes)
        API routes provide authentication profile lookup, complaints, review, assignment, status
        updates, finalization, ratings, comments, attachments, notifications, analytics, activity
        logs, departments, teacher lookup, and weekly reporting. Each protected route validates
        the Supabase session token and checks the user's role before executing business logic.

        4.4.3 Database Technologies
        Supabase PostgreSQL stores profiles, departments, complaints, attachments, comments,
        notifications, category routes, and activity logs. PostgreSQL enums define role, status,
        category, priority, and routed-role values. Supabase Storage is used for complaint
        attachments.
        """,
        fontsize=10.1,
    )
    replace_area(
        doc[35],
        (88, 645, 545, 718),
        """
        15. User logs in through Supabase Auth and receives a session access token.
        16. Client requests include Authorization: Bearer token.
        17. API route validates the token, loads the profile, and checks allowed roles.
        18. Request is rejected with 401/403 when authentication or role scope fails.
        """,
        fontsize=10.1,
    )

    # Page 37: important workflows.
    replace_area(
        doc[36],
        (88, 70, 548, 116),
        """
        19. Complaint queries are scoped by role: Student sees own complaints, Faculty sees
        assigned work, HOD/DSA see department-routed complaints, and Supervisor sees all.
        """,
        fontsize=10.1,
    )
    replace_area(
        doc[36],
        (72, 124, 548, 246),
        """
        4.5.2 AI Category and Priority Suggestion Workflow
        21. Student enters title, description, category, priority, department, and optional
        attachment/anonymous flag.
        22. Keyword-based helpers suggest category and priority before submission.
        23. Category route is selected from category_routes or defaults: Academic, Facilities,
        Behavior-related, and Other route to HOD; Administrative routes to DSA.
        24. Complaint is saved with Submitted status and relevant notifications/activity logs.
        """,
        fontsize=10.1,
    )
    replace_area(
        doc[36],
        (72, 254, 548, 506),
        """
        4.5.3 Severity/Priority Scoring Engine
        27. Define urgency keyword dictionary with weights for urgent, immediately, danger,
        harassment, safety, exam, deadline, broken, failure, critical, problem, issue, angry,
        disappointed, and frustrated.
        28. Convert complaint description to lowercase and remove special characters.
        29. Add keyword weights when words are found in the complaint text.
        30. If score is 7 or above, priority is High.
        31. If score is 4 or above, priority is Medium; otherwise priority is Low.

        4.5.4 Duplicate Awareness / Similarity Check
        32. Convert new and existing complaint descriptions into cleaned token sets.
        33. Compare token overlap using intersection-over-union similarity.
        34. Keep the highest similarity score among existing complaints.
        35. Use similarity insight to help identify repeated or related complaint submissions.
        """,
        fontsize=10.1,
    )
    replace_area(
        doc[36],
        (72, 512, 548, 638),
        """
        4.5.5 HOD/DSA Review and Faculty Assignment Workflow
        38. HOD/DSA opens dashboard and views Submitted complaints within department scope.
        39. HOD/DSA accepts, rejects, or assigns the complaint to a faculty member.
        40. Assignment sets the complaint to In Progress and stores assigned_teacher_id.
        41. Faculty member resolves the complaint with comments; status becomes Resolved.
        42. DSA or Supervisor finalizes the complaint as Closed; student can rate and add
        feedback.
        """,
        fontsize=10.1,
    )
    replace_area(
        doc[36],
        (88, 670, 548, 724),
        """
        44. Analytics and notification APIs identify overdue, rejected, resolved, and pending
        records so staff can follow up quickly from the dashboard.
        """,
        fontsize=10.1,
    )

    # Page 62: conclusion summary.
    replace_area(
        doc[61],
        (88, 140, 548, 314),
        """
        • Full-stack web application with Next.js/React frontend and Next.js API backend.
        • Supabase Auth based authentication with role-based routing and permission checks.
        • Supabase PostgreSQL schema for profiles, departments, complaints, attachments,
        comments, notifications, category routes, and activity logs.
        • Structured complaint lifecycle: Submitted, In Progress, Resolved, Closed, Rejected,
        and Escalated.
        • Dashboards for Student, Faculty Member, HOD, DSA, and Supervisor.
        • Notifications, analytics, weekly reports, chatbot support, and attachment handling.
        """,
        fontsize=10.2,
    )
    replace_area(doc[61], (88, 690, 548, 718), "• Attachment support is implemented through Supabase Storage with allowed file types.")

    # Page 65: references.
    replace_area(
        doc[64],
        (72, 168, 548, 414),
        """
        1. Next.js Documentation. https://nextjs.org/docs
        2. React Documentation. https://react.dev/
        3. Supabase Documentation. https://supabase.com/docs
        4. Supabase Auth Documentation. https://supabase.com/docs/guides/auth
        5. Supabase Storage Documentation. https://supabase.com/docs/guides/storage
        6. PostgreSQL Documentation. https://www.postgresql.org/docs/
        7. Recharts Documentation. https://recharts.org/
        8. Framer Motion Documentation. https://www.framer.com/motion/
        9. Botpress Cloud Documentation. https://botpress.com/docs
        10. Pressman, R.S. (2014). Software Engineering: A Practitioner's Approach (8th ed.).
        11. Sommerville, I. (2015). Software Engineering (10th ed.). Pearson Education.
        """,
        fontsize=10.8,
    )

    # Page 66: stack summary.
    replace_area(
        doc[65],
        (76, 130, 548, 462),
        """
        Frontend               Next.js App Router, React, Framer Motion
        UI/Icons               lucide-react, custom reusable components
        Backend                Next.js API Routes
        Authentication         Supabase Auth session tokens
        Database               Supabase PostgreSQL
        Storage                Supabase Storage complaint-attachments bucket
        Analytics              Recharts and API aggregate endpoints
        Password Security      bcryptjs and password validation helpers
        Chatbot                Botpress Cloud Webchat
        Email/Notifications    Next.js notification/email helper modules
        Version Control        Git / GitHub
        """,
        fontsize=10.8,
    )

    # Page 67: file list.
    replace_area(
        doc[66],
        (72, 126, 548, 452),
        """
        • src/app/api/complaints/route.js — Complaint create/list endpoint.
        • src/app/api/complaints/[id]/review/route.js — HOD/DSA review actions.
        • src/app/api/complaints/[id]/assign/route.js — Faculty assignment workflow.
        • src/app/api/complaints/[id]/status/route.js — Faculty status update workflow.
        • src/app/api/complaints/[id]/finalize/route.js — Complaint closure workflow.
        • src/app/api/complaints/[id]/rate/route.js — Student rating and feedback.
        • src/app/api/complaints/analytics/route.js — Dashboard analytics.
        • src/app/api/complaints/weekly-report/route.js — Weekly report data.
        • src/app/api/departments/route.js — Department management.
        • src/lib/workflow.js — Role scoping and category route defaults.
        • src/lib/auth.js — Supabase session validation and role checks.
        • supabase/schema.sql — PostgreSQL schema, enums, indexes, RLS, and storage bucket.
        """,
        fontsize=10.5,
    )

    doc.save(UPDATED_PDF)
    doc.close()


def pdf_to_docx():
    pdf = fitz.open(UPDATED_PDF)
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)

    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        for index, page in enumerate(pdf):
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image_path = tmp_path / f"page_{index + 1:03}.jpg"
            pix.save(image_path, jpg_quality=92)
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = 0
            paragraph.paragraph_format.space_after = 0
            run = paragraph.add_run()
            run.add_picture(str(image_path), width=Inches(8.5), height=Inches(11))
            if index != len(pdf) - 1:
                document.add_page_break()
    pdf.close()
    document.save(OUT_DOCX)


def main():
    if not SOURCE_PDF.exists():
        raise FileNotFoundError(SOURCE_PDF)
    update_pdf()
    pdf_to_docx()
    print(f"Updated PDF: {UPDATED_PDF}")
    print(f"Updated DOCX: {OUT_DOCX}")


if __name__ == "__main__":
    main()
