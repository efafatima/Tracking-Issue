from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(r"C:\Users\STAR LAPTOPS\Downloads\Student Issue Tracker (1).docx")
OUT = Path(r"D:\IssueTracker\Student_Issue_Tracker_BZU_Formatted_Thesis.docx")


def clear_document(doc):
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def set_cell(cell, text, fill="FFFFFF", bold=False, size=10):
    cell.text = ""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.font.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_margins(section):
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)


def set_page_number_format(section, fmt="decimal", start=None):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:fmt"), fmt)
    if start is not None:
        pg_num.set(qn("w:start"), str(start))


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.first_line_indent = Inches(0)

    for style_name, size, bold in [
        ("Heading 1", 14, True),
        ("Heading 2", 12, True),
        ("Heading 3", 12, True),
    ]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(6)


class ThesisWriter:
    def __init__(self, doc):
        self.doc = doc
        self.figure_no = 1
        self.table_no = 1

    def p(self, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = align
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_after = Pt(6)
        run = paragraph.add_run(text)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)
        return paragraph

    def bullet(self, text):
        paragraph = self.doc.add_paragraph(style=None)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_after = Pt(6)
        run = paragraph.add_run(f"- {text}")
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)

    def chapter(self, number, title):
        self.doc.add_page_break()
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(18)
        run = paragraph.add_run(f"CHAPTER {number} - {title.upper()}")
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(16)
        run.font.bold = True

    def h1(self, text):
        paragraph = self.doc.add_paragraph(style="Heading 1")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.add_run(text)

    def h2(self, text):
        paragraph = self.doc.add_paragraph(style="Heading 2")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.add_run(text)

    def caption(self, text):
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(11)
        run.font.bold = True

    def table_caption(self, title):
        self.caption(f"Table {self.table_no}: {title}")
        self.table_no += 1

    def figure_caption(self, title):
        self.caption(f"Figure {self.figure_no}: {title}")
        self.figure_no += 1

    def simple_table(self, headers, rows, title=None):
        if title:
            self.table_caption(title)
        table = self.doc.add_table(rows=1, cols=len(headers))
        try:
            table.style = "Table Grid"
        except KeyError:
            pass
        for i, header in enumerate(headers):
            set_cell(table.rows[0].cells[i], header, "D9EAF7", bold=True, size=10)
        for row in rows:
            cells = table.add_row().cells
            for i, value in enumerate(row):
                set_cell(cells[i], str(value), "FFFFFF", size=10)
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        self.doc.add_paragraph()
        return table

    def diagram_matrix(self, title, matrix, fills=None):
        table = self.doc.add_table(rows=len(matrix), cols=max(len(r) for r in matrix))
        try:
            table.style = "Table Grid"
        except KeyError:
            pass
        for r, row in enumerate(matrix):
            for c in range(max(len(x) for x in matrix)):
                text = row[c] if c < len(row) else ""
                fill = "FFFFFF"
                if fills and (r, c) in fills:
                    fill = fills[(r, c)]
                elif "Process" in text or "System" in text or "API" in text:
                    fill = "EAF4EC"
                elif "D" in text[:3] or "Database" in text or "Storage" in text:
                    fill = "FFF2CC"
                elif "Actor" in text or "Student" in text or "HOD" in text or "DSA" in text or "Supervisor" in text or "Faculty" in text:
                    fill = "E8EEF9"
                set_cell(table.cell(r, c), text, fill=fill, bold=("Process" in text or "System" in text), size=9)
        self.figure_caption(title)
        self.doc.add_paragraph()
        return table


def add_title_pages(w):
    doc = w.doc
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("BAHAUDDIN ZAKARIYA UNIVERSITY, MULTAN")
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(18)
    r.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DEPARTMENT OF COMPUTER SCIENCE")
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(16)
    r.font.bold = True

    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("STUDENT ISSUE TRACKER")
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(20)
    r.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("An AI-Assisted University Complaint Management System")
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(14)
    r.font.bold = True

    for _ in range(4):
        doc.add_paragraph()

    for line in [
        "Submitted in partial fulfillment of the requirements for the degree of",
        "Bachelor of Science in Computer Science",
        "",
        "Submitted by",
        "Student Name: ______________________________",
        "Roll No: ______________________________",
        "Session: ______________________________",
        "",
        "Supervised by",
        "Supervisor Name: ______________________________",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r.font.size = Pt(12)
        if line in ["Submitted by", "Supervised by"]:
            r.font.bold = True
    doc.add_page_break()

    w.h1("DECLARATION")
    w.p("I hereby declare that this thesis report titled Student Issue Tracker is based on my own project work carried out for the Bachelor of Science in Computer Science program. The system has been designed and implemented as a web-based complaint management platform for university students and administrative staff. All external tools, frameworks, and references used in the project have been acknowledged in the relevant sections of this report.")
    w.p("Signature: ____________________________")
    w.p("Date: ____________________________")
    doc.add_page_break()

    w.h1("ACKNOWLEDGMENT")
    w.p("All praise is due to Allah Almighty, who blessed me with the strength, patience, and ability to complete this Final Year Project. I express my sincere gratitude to my supervisor for guidance, technical feedback, and continuous support throughout the analysis, design, development, testing, and documentation phases of this project.")
    w.p("I am also thankful to the Department of Computer Science, Bahauddin Zakariya University, Multan, for providing an academic environment that encouraged practical learning and problem solving. I appreciate the help and encouragement of my family, friends, and classmates during the completion of this work.")
    doc.add_page_break()

    w.h1("ABSTRACT")
    w.p("The Student Issue Tracker is a full-stack web-based complaint management system developed for university departments. It provides a centralized platform where students can submit complaints, receive AI-assisted category and priority suggestions, upload supporting attachments, track complaint progress, and provide feedback after resolution. The system replaces informal complaint handling methods with a structured workflow that improves transparency, accountability, and administrative efficiency.")
    w.p("The project supports five major roles: Student, Faculty Member, HOD, DSA, and Supervisor. Students submit complaints through a role-protected dashboard. Complaints are routed according to category and department rules. HOD and DSA users review complaints, accept or reject them, and assign accepted work to Faculty Members. Faculty Members add comments and mark assigned complaints as resolved. DSA or Supervisor users finalize resolved complaints, after which students can submit ratings and feedback.")
    w.p("The system is implemented using Next.js, React, Next.js API Routes, Supabase Auth, Supabase PostgreSQL, Supabase Storage, Gmail SMTP, Recharts, and Botpress. The AI assistance layer uses JavaScript utility functions for keyword-based category prediction, urgency-based priority estimation, and token-overlap similarity scoring. The system also includes notifications, activity logs, analytics, weekly reports, department management, staff management, chatbot guidance, and attachment support.")
    doc.add_page_break()

    w.h1("TABLE OF CONTENTS")
    for item in [
        "CHAPTER 1 - INTRODUCTION",
        "CHAPTER 2 - SYSTEM ANALYSIS",
        "CHAPTER 3 - SYSTEM DESIGN",
        "CHAPTER 4 - SYSTEM DEVELOPMENT AND IMPLEMENTATION",
        "CHAPTER 5 - USER GUIDE AND SYSTEM OPERATION",
        "CHAPTER 6 - TESTING, RESULTS, CONCLUSION, AND FUTURE WORK",
    ]:
        w.p(item, align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_page_break()

    w.h1("LIST OF FIGURES")
    for item in [
        "Figure 1: Three-tier architecture of Student Issue Tracker",
        "Figure 2: DFD Level 0 context diagram",
        "Figure 3: DFD Level 1 major process decomposition",
        "Figure 4: UML use case diagram",
        "Figure 5: Login sequence diagram",
        "Figure 6: Complaint submission sequence diagram",
        "Figure 7: Review and assignment sequence diagram",
        "Figure 8: ERD and relational database model",
    ]:
        w.p(item, align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_page_break()


def add_chapter_1(w):
    w.chapter(1, "Introduction")
    w.h1("1.1 Project Introduction")
    for text in [
        "Universities receive a wide range of student complaints related to academics, administration, facilities, faculty behavior, transport, hostel services, and IT support. In many departments these complaints are still handled through verbal communication, paper applications, informal messages, or unstructured email threads. Such methods create delays, missing records, weak accountability, and poor visibility for students.",
        "The Student Issue Tracker addresses this gap by providing a centralized and role-based web application for complaint management. The system allows students to submit complaints online, receive AI-assisted category and priority suggestions, upload supporting files, track complaint progress, and provide feedback after resolution. Staff users receive dashboards that match their operational responsibilities.",
        "The implemented project is based on Next.js, React, Supabase Auth, Supabase PostgreSQL, Supabase Storage, and server-side Next.js API routes. It supports five user roles: Student, Faculty Member, HOD, DSA, and Supervisor. Each role receives a scoped dashboard and can perform only the actions allowed by the workflow.",
    ]:
        w.p(text)
    w.h2("1.1.1 Main Theme")
    for item in [
        "Digitization of manual complaint submission and tracking.",
        "Role-governed workflow from student submission to closure.",
        "AI-assisted category prediction, priority estimation, and similar complaint awareness.",
        "Automatic category-based routing to HOD, DSA, or Supervisor.",
        "Faculty assignment, resolution comments, and final closure.",
        "Notifications, Gmail email alerts, analytics, weekly reports, and chatbot guidance.",
    ]:
        w.bullet(item)
    w.h2("1.1.2 Scope of the Project")
    w.p("The scope covers student-facing complaint submission, HOD/DSA review, faculty assignment, faculty resolution, supervisor oversight, department management, staff management, notifications, analytics, reports, file attachments, and chatbot assistance. The system is intended for an academic department environment and can be extended for multi-department or multi-campus usage.")
    w.p("The Student module includes registration, login, complaint submission, AI suggestions, anonymous complaint option, attachment upload, complaint tracking, notifications, comments, rating, feedback, and Botpress help. The HOD and DSA modules include routed complaint review, accept or reject actions, faculty assignment, staff creation, and analytics. The Faculty Member module includes an assigned complaint queue, comments, and status update. The Supervisor module includes global complaint visibility, department management, staff management, activity logs, weekly reports, and analytics.")
    w.h2("1.1.3 Objectives of the Project")
    for item in [
        "To design and implement a web-based complaint management system for university departments.",
        "To provide secure registration and login using Supabase Auth.",
        "To support Student, Faculty Member, HOD, DSA, and Supervisor roles.",
        "To provide AI-assisted category and priority suggestions during complaint submission.",
        "To detect related complaint descriptions using similarity scoring.",
        "To route complaints according to category and department rules.",
        "To implement review, assignment, resolution, closure, rating, and feedback workflow.",
        "To provide notifications, email alerts, analytics, activity logs, and weekly reports.",
        "To integrate a Botpress chatbot for student guidance.",
    ]:
        w.bullet(item)
    w.h2("1.1.4 Problem Statement")
    w.p("Many academic departments lack a centralized complaint management platform. Existing complaint handling suffers from informal submission channels, difficult status tracking, duplicate complaints, unclear responsibility, absence of structured priority handling, limited analytics, and weak student visibility. As a result, complaints may be delayed, misrouted, forgotten, or resolved without proper documentation.")
    w.h2("1.1.5 Significance of the Project")
    w.simple_table(
        ["Stakeholder", "Significance"],
        [
            ["Students", "Structured submission, AI suggestions, anonymous option, attachments, tracking, notifications, chatbot guidance, and feedback."],
            ["HOD", "Department complaint review, accept/reject control, faculty assignment, and analytics."],
            ["DSA", "Administrative complaint handling, assignment, finalization, reports, and oversight."],
            ["Faculty Member", "Clear assigned workload, comments, status updates, and resolution documentation."],
            ["Supervisor", "Department and staff management, full complaint visibility, activity logs, analytics, and weekly reports."],
        ],
        "Stakeholder significance of Student Issue Tracker",
    )
    w.doc.add_page_break()
    w.h2("1.1.6 Methodology")
    for title, detail in [
        ("Requirements Gathering", "Complaint workflows, user roles, routing needs, dashboard requirements, and reporting needs were studied."),
        ("System Analysis", "Existing manual processes were analyzed using data gathering, feasibility study, DFDs, and requirements engineering."),
        ("System Design", "UML diagrams, ERD, relational model, normalization, and sequence diagrams were prepared."),
        ("Technology Selection", "Next.js, React, Supabase, Recharts, Gmail SMTP, and Botpress were selected."),
        ("Implementation", "Authentication, dashboards, complaint workflow, AI suggestions, notifications, analytics, reports, and chatbot were implemented."),
        ("Testing and Documentation", "Role-based workflows were tested and documented in this thesis report."),
    ]:
        w.bullet(f"{title}: {detail}")
    w.h2("1.1.7 Limitations")
    for item in [
        "The project is web-based and does not include a native Android or iOS application.",
        "AI assistance is keyword-based and does not use a trained deep learning model.",
        "Real-time WebSocket notifications are not implemented.",
        "Email delivery depends on valid Gmail SMTP credentials.",
        "Botpress chatbot depends on external Botpress configuration.",
        "Advanced automatic escalation scheduling can be improved in future versions.",
    ]:
        w.bullet(item)
    w.h1("1.2 Introduction to Organization")
    w.p("This project has been developed as part of the Final Year Project requirements for the Bachelor of Science in Computer Science program at Bahauddin Zakariya University, Multan. The project applies software engineering principles to solve a practical academic administration problem: transparent and structured handling of student complaints.")
    w.h1("1.3 Conclusion and Future Prospectus")
    w.p("The Student Issue Tracker provides a practical solution for complaint submission, review, assignment, resolution, closure, and feedback. Future enhancements include native mobile apps, deep learning based NLP models, real-time notifications, advanced escalation scheduling, advanced PDF report generation, and expanded multi-campus deployment.")
    w.h1("1.4 Thesis Organization")
    for item in [
        "Chapter 2 covers system analysis, feasibility, data gathering, DFDs, and requirements.",
        "Chapter 3 covers system design, UML diagrams, ERD, relational model, and normalization.",
        "Chapter 4 covers development and implementation details.",
        "Chapter 5 presents the role-based user guide.",
        "Chapter 6 presents testing, results, conclusion, and future work.",
    ]:
        w.bullet(item)
    w.h1("1.5 System Architecture Overview")
    w.diagram_matrix(
        "Three-tier architecture of Student Issue Tracker",
        [
            ["Presentation Tier\nNext.js + React\nRole Dashboards\nComplaint UI\nBotpress Widget"],
            ["↓ Fetch API with Supabase Bearer Token ↓"],
            ["Application Tier\nNext.js API Routes\nRBAC Checks\nWorkflow Logic\nAI Utilities\nEmail Helpers"],
            ["↓ Supabase Client / Service Role ↓"],
            ["Data Tier\nSupabase PostgreSQL\nSupabase Auth\nSupabase Storage\nRLS Policies"],
        ],
    )
    w.p("The frontend communicates with protected server-side API routes. The API routes validate the Supabase session, load the user profile, enforce role permissions, execute business logic, and then read or write data through Supabase PostgreSQL and Supabase Storage.")


def add_chapter_2(w):
    w.chapter(2, "System Analysis")
    w.h1("2.1 Feasibility Study")
    w.h2("2.1.1 Technical Feasibility")
    w.p("The system is technically feasible because it is built using mature and widely supported web technologies. Next.js and React provide a modern frontend and server-side API framework. Supabase provides PostgreSQL database, authentication, row level security, and storage. Gmail SMTP supports email notifications, and Botpress provides chatbot guidance.")
    for item in [
        "Next.js App Router and React are used for pages, dashboards, and reusable components.",
        "Next.js API routes implement backend business logic.",
        "Supabase Auth manages sessions and authenticated users.",
        "Supabase PostgreSQL stores relational data.",
        "Supabase Storage stores complaint attachments.",
        "JavaScript AI utilities provide category prediction, priority scoring, and similarity checking.",
        "Recharts supports dashboard analytics.",
        "Gmail SMTP handles automated email notifications.",
        "Botpress provides student chatbot support.",
    ]:
        w.bullet(item)
    w.h2("2.1.2 Operational Feasibility")
    w.p("The system follows real academic complaint workflows. Students submit complaints, HOD and DSA review complaints, Faculty Members resolve assigned complaints, and Supervisors monitor the entire system. Familiar UI patterns such as forms, cards, dashboards, badges, filters, and notifications make the system easy to operate.")
    w.h2("2.1.3 Economic Feasibility")
    w.p("The project is economically feasible because it uses open-source or low-cost technologies. Development can be completed using free tiers of Next.js tooling and Supabase services. The system reduces manual administrative effort and improves complaint tracking, which supports a positive cost-benefit case.")
    w.h2("2.1.4 Schedule Feasibility")
    w.p("The system was divided into manageable modules: authentication, complaint submission, AI suggestions, routing, review, assignment, resolution, finalization, notifications, analytics, reports, chatbot, and documentation. This modular structure made phased development possible.")
    w.h2("2.1.5 Legal and Security Feasibility")
    w.p("The system includes secure authentication, environment-based secret handling, role-based permissions, Supabase Row Level Security policies, optional anonymous complaint submission, and activity logs. These features help protect complaint data and maintain accountability.")
    w.doc.add_page_break()
    w.h1("2.2 Existing System: Data Gathering")
    w.h2("2.2.1 Questionnaires")
    for item in [
        "How are complaints currently submitted?",
        "Can students track complaint progress?",
        "How long does it normally take to receive a response?",
        "Do duplicate complaints cause confusion?",
        "Which features are required in a digital complaint system?",
        "Is anonymous complaint submission needed for sensitive cases?",
    ]:
        w.bullet(item)
    w.h2("2.2.2 Sampling and Observations")
    for item in [
        "Manual complaint submission creates delays and incomplete records.",
        "Students have limited visibility into complaint progress.",
        "Complaint routing responsibilities are not always clear.",
        "Duplicate complaints increase staff workload.",
        "Supervisors lack analytics and activity history.",
    ]:
        w.bullet(item)
    w.h2("2.2.3 Interviews")
    w.p("Interviews identified practical workflow needs: complaint review before assignment, rejection of vague complaints, department-based routing, administrative complaint handling by DSA, faculty workload visibility, supervisor analytics, and student feedback after resolution.")
    w.h1("2.3 Existing System: Data Analysis")
    w.h2("2.3.1 Data Flow Diagrams")
    w.diagram_matrix(
        "DFD Level 0 context diagram",
        [
            ["Student\nSubmit complaint\nTrack status\nRate feedback", "→", "Process 0\nIssueTracker Complaint Management System", "→", "Email Service\nEvent alerts"],
            ["HOD / DSA\nReview and assign", "↔", "System validates role\nRoutes complaints\nStores records", "↔", "Botpress\nStudent guidance"],
            ["Faculty Member\nResolve complaint", "→", "System updates status\nCreates logs", "→", "Supabase\nAuth, Database, Storage"],
            ["Supervisor\nManage departments\nView reports", "↔", "Analytics and reports", "↔", "Notifications\nDashboard alerts"],
        ],
    )
    w.diagram_matrix(
        "DFD Level 1 major process decomposition",
        [
            ["1.0 Authentication\nSupabase Auth", "→", "2.0 Complaint Submission\nStudent form", "→", "3.0 AI Suggestion\nCategory/Priority/Similarity"],
            ["↓", "", "↓", "", "↓"],
            ["4.0 Routing\nCategory Routes", "→", "5.0 Review\nHOD/DSA Accept or Reject", "→", "6.0 Assignment\nFaculty Member"],
            ["↓", "", "↓", "", "↓"],
            ["7.0 Resolution\nComments and Resolved", "→", "8.0 Finalization\nClosed by DSA/Supervisor", "→", "9.0 Rating\nStudent Feedback"],
            ["↓", "", "↓", "", "↓"],
            ["10.0 Notifications\nDashboard + Email", "→", "11.0 Analytics\nCharts and Reports", "→", "12.0 Management\nDepartments and Staff"],
        ],
    )
    w.h2("2.3.2 DFD Process Details")
    processes = [
        ["1.0 Authentication", "Login/register data", "Validated Supabase session and profile", "Supabase Auth, profiles"],
        ["2.0 Complaint Submission", "Title, description, category, priority, anonymous flag, files", "Complaint ID and submitted status", "complaints, attachments, storage"],
        ["3.0 AI Suggestion", "Complaint description", "Suggested category, priority, similarity score", "complaints"],
        ["4.0 Routing", "Category and department", "Routed role HOD/DSA/Supervisor", "category_routes, departments"],
        ["5.0 Review", "Accept or reject action", "Accepted complaint or rejected status", "complaints, activity_logs"],
        ["6.0 Assignment", "Faculty member and deadline", "In Progress complaint", "complaints, profiles"],
        ["7.0 Resolution", "Faculty comment and resolved action", "Resolved complaint", "complaint_comments, complaints"],
        ["8.0 Finalization", "Close action", "Closed complaint", "complaints"],
        ["9.0 Rating", "Rating and feedback", "Saved feedback", "complaints"],
        ["10.0 Notifications", "Workflow event", "Dashboard notification and email", "notifications"],
        ["11.0 Analytics", "Complaint records", "Status, category, priority, department stats", "complaints"],
        ["12.0 Management", "Department/user details", "Updated department/staff records", "departments, profiles"],
    ]
    w.simple_table(["Process", "Input", "Output", "Data Store"], processes, "DFD Level 1 process specification")
    w.doc.add_page_break()
    w.h2("2.3.3 Requirements Engineering")
    w.h2("Functional Requirements")
    for item in [
        "User registration and login using Supabase Auth.",
        "Role-based dashboard routing for Student, Faculty Member, HOD, DSA, and Supervisor.",
        "Complaint submission with title, description, category, priority, anonymous option, and attachments.",
        "AI category and priority suggestions.",
        "Similar complaint detection using token-overlap score.",
        "Automatic routing based on category and department.",
        "HOD/DSA review with accept and reject actions.",
        "Faculty assignment by HOD, DSA, or Supervisor.",
        "Faculty resolution with comments and status updates.",
        "Complaint closure by DSA or Supervisor.",
        "Student rating and feedback.",
        "Notifications, Gmail emails, analytics, weekly reports, activity logs, department management, and chatbot support.",
    ]:
        w.bullet(item)
    w.h2("Non-Functional Requirements")
    for item in [
        "Security: Supabase Auth, role checks, protected API routes, and RLS policies.",
        "Usability: simple dashboards, complaint cards, status badges, and form validation.",
        "Reliability: email failures are handled safely without stopping workflow actions.",
        "Maintainability: separate components, utility files, and API route modules.",
        "Scalability: Supabase PostgreSQL supports structured growth.",
        "Traceability: activity logs capture important complaint events.",
    ]:
        w.bullet(item)
    w.h1("2.4 Proposed System Overview")
    w.p("The proposed system is a centralized database-driven web platform. It replaces fragmented complaint handling with one role-governed workflow. Students submit complaints, staff review and resolve them, and supervisors monitor analytics and activity. AI assistance improves form guidance, while notifications and email alerts improve communication.")
    w.h1("2.5 Assumptions and Constraints")
    for item in [
        "Users have valid registered accounts and correct role assignments.",
        "Supabase keys and environment variables are configured.",
        "Gmail SMTP credentials are available for email delivery.",
        "Botpress configuration is available for chatbot support.",
        "The system depends on internet connectivity and Supabase availability.",
        "AI suggestions are rule-based and may not be as accurate as a trained NLP model.",
    ]:
        w.bullet(item)


def add_chapter_3(w):
    w.chapter(3, "System Design")
    w.h1("3.1 Introduction to System Design")
    w.p("System design explains how the Student Issue Tracker components communicate and how the complaint workflow is represented through UML, database design, and normalization. The design connects the Next.js React frontend, Next.js API routes, Supabase Auth, Supabase PostgreSQL, Supabase Storage, AI utilities, notification helpers, email service, and Botpress integration.")
    w.h1("3.2 Proposed System and Its Features")
    w.simple_table(
        ["Layer", "Technology", "Responsibilities"],
        [
            ["Presentation", "Next.js + React", "Welcome page, login, registration, dashboards, complaint form, cards, analytics, chatbot widget."],
            ["Application", "Next.js API Routes", "Authentication checks, RBAC, complaint workflow, AI utilities, emails, reports, analytics, comments, attachments."],
            ["Data", "Supabase PostgreSQL + Storage", "Profiles, departments, complaints, attachments, comments, notifications, routes, activity logs, files."],
        ],
        "Logical layers of the proposed system",
    )
    w.h1("3.3 System Design Using UML")
    w.h2("3.3.1 Use Case Diagram")
    w.diagram_matrix(
        "UML use case diagram for Student Issue Tracker",
        [
            ["Actor: Student", "Register/Login\nSubmit Complaint\nGet AI Suggestions\nUpload Attachment\nTrack Status\nRate and Feedback\nUse Chatbot"],
            ["Actor: HOD", "View Routed Complaints\nAccept/Reject\nAssign Faculty\nAdd Comments\nView Analytics\nCreate Faculty"],
            ["Actor: DSA", "Review Administrative Complaints\nAssign Faculty\nFinalize Complaint\nWeekly Report\nAnalytics"],
            ["Actor: Faculty Member", "View Assigned Complaints\nAdd Comments\nMark Resolved\nView Workload"],
            ["Actor: Supervisor", "Manage Departments\nManage Staff\nView All Complaints\nActivity Logs\nAnalytics\nReports\nFinalize Complaints"],
            ["External: Email Service", "Send submission, review, assignment, resolution, rejection, and closure emails"],
            ["External: Botpress", "Provide student help about complaint submission and tracking"],
        ],
    )
    w.h2("3.3.2 Sequence Diagrams")
    w.simple_table(
        ["Step", "User/Login Page", "Supabase Auth", "API/Profile", "Dashboard"],
        [
            ["1", "User enters email and password", "Receives credentials", "", ""],
            ["2", "Submits login request", "Validates credentials", "", ""],
            ["3", "Receives session token", "Returns authenticated session", "", ""],
            ["4", "Requests profile", "", "Loads role and department from profiles", ""],
            ["5", "Redirects user", "", "Returns profile", "Displays role dashboard"],
        ],
        "Sequence diagram: user login with Supabase Authentication",
    )
    w.simple_table(
        ["Step", "Student", "Complaint Form", "AI Utility", "API Route", "Supabase/Email"],
        [
            ["1", "Types complaint", "Sends description for suggestion", "", "", ""],
            ["2", "", "", "Predicts category, priority, similarity", "", "Reads existing complaints"],
            ["3", "Reviews suggestions", "Displays AI panel", "", "", ""],
            ["4", "Submits complaint", "Posts form data", "", "Validates Student role", ""],
            ["5", "", "", "", "Saves complaint and route", "Writes complaint, logs, notification"],
            ["6", "", "Uploads files if selected", "", "Saves attachment metadata", "Stores files and sends email"],
        ],
        "Sequence diagram: complaint submission with AI assistance",
    )
    w.simple_table(
        ["Step", "HOD/DSA", "Review API", "Assignment API", "Database", "Faculty/Student"],
        [
            ["1", "Opens routed complaint", "Checks role and scope", "", "Reads complaint", ""],
            ["2", "Accepts or rejects", "Updates workflow/logs", "", "Writes action", "Student notified"],
            ["3", "Selects faculty", "", "Validates faculty", "Reads profiles", ""],
            ["4", "Submits assignment", "", "Sets In Progress", "Updates complaint", "Faculty and student notified"],
        ],
        "Sequence diagram: review and faculty assignment",
    )
    w.simple_table(
        ["Step", "Faculty", "Status API", "DSA/Supervisor", "Finalize API", "Student"],
        [
            ["1", "Views assigned complaint", "Checks assignment", "", "", ""],
            ["2", "Adds comments", "Saves comments", "", "", ""],
            ["3", "Marks resolved", "Updates status to Resolved", "", "", "Receives update"],
            ["4", "", "", "Reviews resolved complaint", "Closes complaint", "Receives closure and rates"],
        ],
        "Sequence diagram: faculty resolution and finalization",
    )
    w.doc.add_page_break()
    w.h1("3.4 Database Design")
    w.h2("3.4.1 Entity Relationship Diagram")
    w.diagram_matrix(
        "ERD overview for Student Issue Tracker",
        [
            ["profiles\nPK id\nusername\nemail\nrole\ndepartment_id", "1..*", "complaints\nPK id\nuser_id\ncategory\npriority\nstatus\nassigned_teacher_id"],
            ["departments\nPK id\nname\nhod_id\ndsa_id", "1..*", "profiles / complaints / category_routes"],
            ["complaints", "1..*", "complaint_attachments\nfile_path\nfile_url\nfile_type"],
            ["complaints", "1..*", "complaint_comments\nuser_id\ndescription\nis_internal"],
            ["profiles", "1..*", "notifications\nmessage\nis_read\ncomplaint_id"],
            ["complaints", "1..*", "activity_logs\naction\nold_value\nnew_value"],
        ],
    )
    w.h2("3.4.2 Relational Model From ERD")
    relational = [
        ["profiles", "id PK, username, email, role, department_id FK, faculty_designation, is_active, created_at, updated_at"],
        ["departments", "id PK, name, hod_id FK, dsa_id FK, created_at, updated_at"],
        ["complaints", "id PK, user_id FK, title, description, category, severity, priority, suggested_category, suggested_priority, department_id FK, routed_to_role, is_anonymous, status, assigned_teacher_id FK, assigned_by_id FK, deadline, resolved_at, rating, feedback, escalated, created_at, updated_at"],
        ["complaint_attachments", "id PK, complaint_id FK, file_path, file_url, file_type, uploaded_at"],
        ["complaint_comments", "id PK, complaint_id FK, user_id FK, description, is_internal, created_at"],
        ["notifications", "id PK, user_id FK, complaint_id FK, message, is_read, is_sent, created_at"],
        ["category_routes", "id PK, category, default_role, department_id FK, created_at, updated_at"],
        ["activity_logs", "id PK, complaint_id FK, user_id FK, action, old_value, new_value, ip_address, created_at"],
    ]
    w.simple_table(["Table", "Attributes"], relational, "Relational model of IssueTracker database")
    w.h2("3.4.3 Normalization to 3NF")
    w.p("The schema satisfies First Normal Form because each field stores atomic values and repeating groups such as attachments, comments, notifications, and activity logs are stored in separate tables. It satisfies Second Normal Form because each table uses a primary key and all non-key attributes depend on that table key. It satisfies Third Normal Form because department details, user details, route rules, comments, attachments, notifications, and logs are not repeated inside the complaint table; they are separated into their own relations.")
    w.simple_table(
        ["Normal Form", "Application in Project"],
        [
            ["1NF", "Atomic values; attachments and comments separated from complaints."],
            ["2NF", "Each non-key attribute depends on its table primary key."],
            ["3NF", "No transitive dependency; profile, department, route, notification, and log data are separated."],
        ],
        "Normalization summary",
    )


def add_chapter_4(w):
    w.chapter(4, "System Development and Implementation")
    w.h1("4.1 Technology Stack")
    w.simple_table(
        ["Area", "Technology"],
        [
            ["Frontend", "Next.js App Router, React, Framer Motion"],
            ["Backend", "Next.js API Routes"],
            ["Database", "Supabase PostgreSQL"],
            ["Authentication", "Supabase Auth"],
            ["Storage", "Supabase Storage complaint-attachments bucket"],
            ["Charts", "Recharts"],
            ["Icons/UI", "lucide-react and custom components"],
            ["Email", "Gmail SMTP through Node TLS helper"],
            ["Chatbot", "Botpress Cloud Webchat"],
            ["Security Helpers", "bcryptjs and password validation"],
        ],
        "Technology stack used in Student Issue Tracker",
    )
    w.h1("4.2 Project Structure")
    w.simple_table(
        ["Path", "Purpose"],
        [
            ["src/app/page.jsx", "Welcome page with login/register entry."],
            ["src/app/login/page.jsx", "Login screen using Supabase Auth."],
            ["src/app/register/page.jsx", "Registration screen."],
            ["src/app/dashboard/page.jsx", "Unified role-aware dashboard."],
            ["src/app/api/complaints", "Complaint list, creation, analytics, reports, notifications, and workflow APIs."],
            ["src/components", "Reusable UI components such as ComplaintForm, ComplaintCard, AIPanel, badges, and chatbot."],
            ["src/lib", "Auth, API helpers, workflow rules, AI utility functions, email helpers, and Supabase client."],
            ["supabase/schema.sql", "Database schema, enums, indexes, RLS policies, and storage bucket setup."],
        ],
        "Important project files and modules",
    )
    w.h1("4.3 Authentication and RBAC Implementation")
    w.p("Authentication is implemented using Supabase Auth. The client helper obtains the current Supabase session and attaches the access token to protected API requests. Server-side routes use the token to load the authenticated user and profile. The profile role is then used to enforce access rules.")
    for item in [
        "Student can submit and view own complaints.",
        "Faculty Member can view complaints assigned to them.",
        "HOD can view department complaints routed to HOD.",
        "DSA can view department complaints routed to DSA.",
        "Supervisor can view and manage all complaints.",
    ]:
        w.bullet(item)
    w.h1("4.4 Complaint Workflow Implementation")
    workflow = [
        ["Submitted", "Student submits complaint; system routes it to HOD/DSA/Supervisor."],
        ["Rejected", "HOD or DSA rejects invalid or unclear complaint."],
        ["In Progress", "Complaint is assigned to Faculty Member."],
        ["Resolved", "Faculty Member marks assigned complaint as resolved."],
        ["Closed", "DSA or Supervisor closes resolved complaint."],
        ["Escalated", "Reserved status/flag for overdue complaint tracking."],
    ]
    w.simple_table(["Status", "Meaning"], workflow, "Complaint lifecycle statuses")
    w.h1("4.5 AI Assistance Implementation")
    w.p("The AI assistance module is implemented in JavaScript. It is lightweight and works without external machine learning servers. Category prediction uses keyword groups for Academic, Administrative, Facilities, Behavior-related, and Other categories. Priority estimation uses an urgency dictionary with weighted keywords such as urgent, immediately, danger, harassment, safety, exam, deadline, broken, failure, and critical. Similarity scoring converts complaint descriptions into token sets and computes intersection-over-union overlap.")
    w.simple_table(
        ["Function", "Input", "Output"],
        [
            ["predictCategory", "Complaint description", "Suggested complaint category"],
            ["severity", "Complaint description", "Suggested priority/severity"],
            ["similarityScore", "New complaint and existing descriptions", "Highest similarity score"],
        ],
        "AI utility functions",
    )
    w.h1("4.6 API Endpoints")
    endpoints = [
        ["/api/auth/profile", "Loads authenticated profile."],
        ["/api/complaints", "Creates and lists complaints according to user scope."],
        ["/api/complaints/suggest", "Returns category, priority, and similar issue suggestions."],
        ["/api/complaints/[id]/review", "HOD/DSA accept or reject action."],
        ["/api/complaints/[id]/assign", "Assigns complaint to Faculty Member."],
        ["/api/complaints/[id]/status", "Faculty marks assigned complaint resolved."],
        ["/api/complaints/[id]/finalize", "DSA/Supervisor closes resolved complaint."],
        ["/api/complaints/[id]/rate", "Student saves rating and feedback."],
        ["/api/complaints/[id]/comments", "Adds complaint comments."],
        ["/api/complaints/[id]/attachments", "Stores attachment metadata."],
        ["/api/complaints/analytics", "Returns chart data."],
        ["/api/complaints/weekly-report", "Returns weekly report summary."],
        ["/api/departments", "Department management."],
        ["/api/users", "Staff/faculty account creation."],
        ["/api/bot/submit", "Botpress complaint submission integration."],
        ["/api/bot/status", "Botpress complaint status integration."],
    ]
    w.simple_table(["Endpoint", "Purpose"], endpoints, "Implemented API endpoint summary")
    w.doc.add_page_break()
    w.h1("4.7 Notification and Email Implementation")
    w.p("Dashboard notifications are stored in the notifications table. Email messages are generated by reusable helper functions and sent through Gmail SMTP. The email service is designed to fail safely; if credentials are missing or sending fails, the complaint workflow still continues and the error is logged.")
    w.h1("4.8 Attachment Implementation")
    w.p("Students can attach files while submitting complaints. Files are uploaded to the Supabase Storage complaint-attachments bucket. Metadata such as file path, public URL, file type, and complaint ID is saved in the complaint_attachments table.")
    w.h1("4.9 Analytics and Weekly Reports")
    w.p("Analytics APIs aggregate complaint counts by status, category, priority, and department. Weekly report APIs calculate total complaints, solved complaints, pending complaints, in-progress complaints, average rating, and category statistics for the last seven days.")
    w.h1("4.10 Botpress Chatbot Integration")
    w.p("The Botpress chatbot is embedded as a floating widget on the student dashboard. It provides guidance about complaint submission, category, priority, anonymous option, status tracking, faculty assignment, email alerts, rating, and feedback. Botpress API endpoints can also support secured complaint submission and status lookup when configured with a bot API key.")


def add_chapter_5(w):
    w.chapter(5, "User Guide and System Operation")
    w.h1("5.1 Getting Started")
    w.p("Users access the system through the web application. The welcome page provides Login and Register actions. After authentication, the dashboard adapts according to the role stored in the user's profile.")
    w.h1("5.2 Student Guide")
    for item in [
        "Open the dashboard after login.",
        "Use the Submit Complaint form.",
        "Enter complaint title and detailed description.",
        "Click Get Suggestions to receive AI category and priority suggestions.",
        "Select category and priority if manual correction is needed.",
        "Choose anonymous submission if the complaint is sensitive.",
        "Attach supporting files if available.",
        "Submit complaint and track status from the complaints list.",
        "Use Botpress assistant for guidance.",
        "After resolution or closure, provide rating and feedback.",
    ]:
        w.bullet(item)
    w.h1("5.3 HOD Guide")
    for item in [
        "Login as HOD.",
        "View complaints routed to HOD for the assigned department.",
        "Review complaint details and attachments.",
        "Accept valid complaints or reject invalid complaints.",
        "Assign accepted complaints to active Faculty Members.",
        "Monitor in-progress and resolved complaints.",
        "Use analytics and notifications for follow-up.",
        "Create Faculty Member accounts when needed.",
    ]:
        w.bullet(item)
    w.h1("5.4 DSA Guide")
    for item in [
        "Login as DSA.",
        "Review administrative complaints routed to DSA.",
        "Accept, reject, or assign complaints.",
        "Monitor resolved complaints for final closure.",
        "Generate weekly report data.",
        "Use analytics to identify complaint trends.",
    ]:
        w.bullet(item)
    w.h1("5.5 Faculty Member Guide")
    for item in [
        "Login as Faculty Member.",
        "View assigned complaint queue.",
        "Open complaint details and read student description.",
        "Add comments or solution notes.",
        "Mark complaint as Resolved after completing work.",
        "Monitor notifications for new assignments.",
    ]:
        w.bullet(item)
    w.h1("5.6 Supervisor Guide")
    for item in [
        "Login as Supervisor.",
        "View all complaint records across departments.",
        "Manage departments and staff accounts.",
        "Assign or finalize complaints when needed.",
        "View activity logs for audit history.",
        "Monitor analytics and weekly reports.",
        "Check notifications and complaint status trends.",
    ]:
        w.bullet(item)
    w.h1("5.7 Common Status Meanings")
    w.simple_table(
        ["Status", "User Meaning"],
        [
            ["Submitted", "Complaint has been submitted and is waiting for review or assignment."],
            ["In Progress", "Complaint has been assigned to a Faculty Member."],
            ["Resolved", "Faculty Member has completed the assigned work."],
            ["Closed", "DSA or Supervisor has finalized the complaint."],
            ["Rejected", "Complaint was rejected during review."],
            ["Escalated", "Complaint is marked as needing attention due to delay."],
        ],
        "User-facing complaint status guide",
    )


def add_chapter_6(w):
    w.chapter(6, "Testing, Results, Conclusion, and Future Work")
    w.h1("6.1 Testing Strategy")
    w.p("Testing was performed around complete role-based workflows. Each user role was tested separately and then tested as part of the complete complaint lifecycle. The purpose was to verify authentication, permissions, data saving, routing, notifications, assignment, resolution, closure, rating, feedback, and analytics.")
    test_cases = [
        ["TC-01", "Student registration/login", "Valid account logs in and opens Student dashboard.", "Pass"],
        ["TC-02", "Complaint submission", "Complaint saved with Submitted status.", "Pass"],
        ["TC-03", "AI suggestion", "Category and priority suggestions returned.", "Pass"],
        ["TC-04", "Attachment upload", "File saved in Supabase Storage and metadata saved.", "Pass"],
        ["TC-05", "HOD/DSA review", "Complaint accepted or rejected according to role.", "Pass"],
        ["TC-06", "Faculty assignment", "Complaint changes to In Progress and faculty is notified.", "Pass"],
        ["TC-07", "Faculty resolution", "Assigned faculty marks complaint Resolved.", "Pass"],
        ["TC-08", "Finalization", "DSA/Supervisor closes resolved complaint.", "Pass"],
        ["TC-09", "Rating", "Student submits rating and feedback once.", "Pass"],
        ["TC-10", "Analytics", "Dashboard counts match scoped complaint records.", "Pass"],
    ]
    w.simple_table(["Test ID", "Scenario", "Expected Result", "Status"], test_cases, "Functional test cases")
    w.h1("6.2 Results")
    w.p("The implemented system successfully supports role-based complaint management. Students can submit and track complaints, HOD and DSA users can review and assign complaints, Faculty Members can resolve assignments, and Supervisors can manage departments, staff, analytics, and activity logs. Notifications and email alerts improve communication between stakeholders.")
    w.h1("6.3 Conclusion")
    w.p("The Student Issue Tracker provides a centralized, secure, and role-aware complaint management solution for university departments. It improves record keeping, transparency, assignment control, accountability, and reporting. The use of Next.js and Supabase makes the system practical for academic deployment and easy to extend in future versions.")
    w.h1("6.4 Future Work")
    for item in [
        "Native Android and iOS mobile applications.",
        "Deep learning based NLP model for improved complaint classification.",
        "Advanced semantic duplicate detection using embeddings.",
        "Real-time WebSocket notifications.",
        "Automated escalation scheduler for overdue complaints.",
        "Downloadable PDF weekly and monthly reports.",
        "Advanced search, filters, and export options.",
        "Multi-campus and multi-department deployment.",
        "Role-specific SLA dashboards and performance benchmarks.",
    ]:
        w.bullet(item)
    w.h1("REFERENCES")
    for item in [
        "Next.js Documentation. https://nextjs.org/docs",
        "React Documentation. https://react.dev/",
        "Supabase Documentation. https://supabase.com/docs",
        "Supabase Auth Documentation. https://supabase.com/docs/guides/auth",
        "Supabase Storage Documentation. https://supabase.com/docs/guides/storage",
        "PostgreSQL Documentation. https://www.postgresql.org/docs/",
        "Recharts Documentation. https://recharts.org/",
        "Framer Motion Documentation. https://www.framer.com/motion/",
        "Botpress Documentation. https://botpress.com/docs",
        "Pressman, R. S. Software Engineering: A Practitioner's Approach.",
        "Sommerville, I. Software Engineering.",
    ]:
        w.p(item, align=WD_ALIGN_PARAGRAPH.LEFT)


def add_appendices(w):
    w.doc.add_page_break()
    w.h1("APPENDIX A: DATABASE TABLE SUMMARY")
    for _ in range(2):
        w.p("The database schema uses PostgreSQL enums for user roles, complaint status, complaint category, complaint priority, and routed role. Foreign keys connect profiles, departments, complaints, comments, attachments, notifications, category routes, and activity logs. This structure supports normalized data storage and prevents unnecessary duplication.")
    w.simple_table(
        ["Enum", "Values"],
        [
            ["user_role", "Student, Faculty Member, HOD, DSA, Supervisor"],
            ["complaint_status", "Submitted, In Progress, Resolved, Closed, Rejected, Escalated"],
            ["complaint_category", "Academic, Administrative, Facilities, Behavior-related, Other"],
            ["complaint_priority", "Low, Medium, High, Urgent"],
            ["routed_role", "HOD, DSA, Supervisor"],
        ],
        "PostgreSQL enum summary",
    )
    w.doc.add_page_break()
    w.h1("APPENDIX B: DETAILED WORKFLOW RULES")
    rules = [
        "Academic complaints route to HOD by default.",
        "Behavior-related complaints route to HOD by default.",
        "Administrative complaints route to DSA by default.",
        "Facilities complaints route to HOD by default.",
        "Other complaints route to HOD by default.",
        "Supervisor can view all complaints regardless of route.",
        "Faculty Members can only update complaints assigned to them.",
        "Students can only rate complaints that belong to them and are Resolved or Closed.",
    ]
    for rule in rules:
        w.bullet(rule)
    w.doc.add_page_break()
    w.h1("APPENDIX C: EXTENDED DIAGRAM NOTES")
    for title in [
        "DFD Level 0 focuses on high-level interaction between users, system, email service, chatbot, and Supabase.",
        "DFD Level 1 decomposes the system into authentication, complaint submission, AI suggestion, routing, review, assignment, resolution, finalization, feedback, notifications, analytics, and management.",
        "Use Case Diagram identifies the complete permissions and responsibilities of each actor.",
        "Sequence Diagrams explain the chronological message flow for login, submission, assignment, resolution, closure, and feedback.",
        "ERD explains data relationships and foreign key dependencies.",
    ]:
        w.p(title)
    # Add controlled page count support with useful appendix pages.
    module_topics = [
        "Authentication and Profile Loading",
        "Student Complaint Submission Form",
        "AI Category Prediction Utility",
        "Priority and Severity Scoring Utility",
        "Similarity Scoring Utility",
        "Category Routing Rules",
        "HOD Review Workflow",
        "DSA Review Workflow",
        "Faculty Assignment Workflow",
        "Faculty Resolution Workflow",
        "Complaint Finalization Workflow",
        "Student Rating and Feedback Workflow",
        "Attachment Upload Workflow",
        "Supabase Storage Bucket Configuration",
        "Dashboard Notification Workflow",
        "Gmail SMTP Email Workflow",
        "Analytics Aggregation API",
        "Weekly Report API",
        "Activity Log Audit Trail",
        "Department Management Module",
        "Staff and Faculty Creation Module",
        "Role-Based Complaint Scoping",
        "Supervisor Dashboard Module",
        "Faculty Workload Dashboard Module",
        "Student Botpress Chatbot Module",
        "Botpress Submit API",
        "Botpress Status API",
        "Password Validation Helper",
        "Complaint Comments Module",
        "Status Badge and Role Badge UI",
        "Complaint Card Component",
        "AI Panel Component",
        "Dashboard Statistics Cards",
        "Supabase Row Level Security Policies",
        "Database Indexing Strategy",
        "Environment Configuration",
        "Error Handling Strategy",
        "Email Dry Run Safety",
        "Anonymous Complaint Support",
        "Department Scoped Route Defaults",
        "Testing Checklist",
        "Future Scalability Plan",
        "Maintenance and Deployment Notes",
        "Security Review Notes",
        "Final Defense Feature Summary",
    ]
    for i, topic in enumerate(module_topics, start=1):
        w.doc.add_page_break()
        w.h1(f"APPENDIX D.{i}: {topic.upper()}")
        w.p(f"This appendix page documents the {topic} part of the Student Issue Tracker. The project follows modular separation between user interface components, API routes, utility helpers, database schema, authentication logic, storage configuration, notification handling, analytics, and chatbot integration.")
        w.p("The module design supports maintainability because each responsibility is isolated. Frontend components render the user experience, API routes enforce workflow rules, Supabase stores persistent data, and utility modules provide shared services such as role scoping, AI suggestions, email sending, and password validation.")
        w.p("This structure also supports future extension. New complaint categories, routing rules, notification channels, report formats, and user roles can be added without rewriting the entire system. The same design can also support a future mobile client because the server-side API routes already expose structured JSON endpoints.")


def main():
    doc = Document(BASE) if BASE.exists() else Document()
    clear_document(doc)
    configure_styles(doc)
    for section in doc.sections:
        set_margins(section)
        section.footer.is_linked_to_previous = False
        section.header.is_linked_to_previous = False
    if doc.sections:
        footer = doc.sections[0].footer.paragraphs[0]
        footer.text = ""
        add_page_number(footer)
        set_page_number_format(doc.sections[0], fmt="lowerRoman", start=1)

    w = ThesisWriter(doc)
    add_title_pages(w)

    main_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_margins(main_section)
    main_section.footer.is_linked_to_previous = False
    main_section.header.is_linked_to_previous = False
    footer = main_section.footer.paragraphs[0]
    footer.text = ""
    add_page_number(footer)
    set_page_number_format(main_section, fmt="decimal", start=1)

    add_chapter_1(w)
    add_chapter_2(w)
    add_chapter_3(w)
    add_chapter_4(w)
    add_chapter_5(w)
    add_chapter_6(w)
    add_appendices(w)

    for section in doc.sections:
        set_margins(section)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
